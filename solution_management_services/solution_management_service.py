"""方案管理服务子模块

负责对方案文档对象进行增、删、改、查、存，
并转化为纯文本格式以备用。
"""
import os
import json
import logging
from datetime import datetime
from typing import Optional, Dict, List, Any
from pydantic import BaseModel, Field

import sys
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))) + '/..')
from bo.solution import SolutionDocument, DocumentType


class DocumentQueryFilter(BaseModel):
    """文档查询过滤条件"""
    document_type: Optional[DocumentType] = Field(default=None, description="文档类型")
    file_format: Optional[str] = Field(default=None, description="文件格式")
    created_by: Optional[str] = Field(default=None, description="创建人")
    keyword: Optional[str] = Field(default=None, description="关键词（模糊匹配文件名和内容）")


class SolutionManagementService:
    """
    方案管理服务子模块
    
    提供对方案文档对象的完整生命周期管理：
    - 创建（Create）：创建新的方案文档
    - 查询（Read）：根据条件查询文档
    - 更新（Update）：修改文档内容和属性
    - 删除（Delete）：删除文档
    - 存储（Save）：持久化文档到存储介质
    - 转换（Convert）：将文档转换为纯文本格式
    
    :param storage_path: 文档存储路径
    :param logger: 日志记录器
    """
    
    def __init__(self, storage_path: str = "documents", logger: Optional[logging.Logger] = None):
        self._storage_path = storage_path
        self._documents: Dict[str, SolutionDocument] = {}
        self._logger = logger or self._setup_logging()
        
        # 确保存储目录存在
        os.makedirs(self._storage_path, exist_ok=True)
        self._logger.info(f"方案管理服务已初始化，存储路径: {self._storage_path}")
    
    def _setup_logging(self) -> logging.Logger:
        """配置日志系统"""
        logger = logging.getLogger("SolutionManagementService")
        logger.setLevel(logging.INFO)
        
        if not logger.handlers:
            handler = logging.StreamHandler()
            formatter = logging.Formatter(
                '%(asctime)s - %(name)s - %(levelname)s - %(message)s'
            )
            handler.setFormatter(formatter)
            logger.addHandler(handler)
        
        return logger
    
    # ========== 创建操作 ==========
    
    def create_document(self, document: SolutionDocument) -> SolutionDocument:
        """
        创建方案文档
        
        :param document: 方案文档对象
        :return: 创建后的文档对象
        :raises ValueError: 如果文档ID已存在
        """
        if document.document_id in self._documents:
            raise ValueError(f"文档ID {document.document_id} 已存在")
        
        self._documents[document.document_id] = document
        self._logger.info(f"创建文档: {document.document_id} - {document.file_name}")
        return document
    
    def create_document_from_file(
        self,
        document_id: str,
        file_path: str,
        version: str = "1.0",
        document_type: DocumentType = DocumentType.MAIN,
        description: Optional[str] = None
    ) -> SolutionDocument:
        """
        从文件创建方案文档
        
        :param document_id: 文档唯一标识
        :param file_path: 文件路径
        :param version: 版本号
        :param document_type: 文档类型
        :param description: 文档描述
        :return: 创建的文档对象
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        file_name = os.path.basename(file_path)
        file_format = os.path.splitext(file_name)[1].lstrip('.')
        file_size = os.path.getsize(file_path)
        
        # 读取文件内容
        try:
            with open(file_path, 'rb') as f:
                file_content = f.read()
        except Exception as e:
            self._logger.error(f"读取文件失败: {file_path}, 错误: {e}")
            raise
        
        # 尝试提取文本内容
        text_content = self._extract_text_from_file(file_path, file_content)
        
        document = SolutionDocument(
            document_id=document_id,
            file_name=file_name,
            version=version,
            document_type=document_type,
            file_content=file_content,
            text_content=text_content,
            description=description,
            format=file_format,
            size=file_size
        )
        
        return self.create_document(document)
    
    def _extract_text_from_file(self, file_path: str, file_content: bytes) -> Optional[str]:
        """
        从文件中提取文本内容
        
        :param file_path: 文件路径
        :param file_content: 文件二进制内容
        :return: 提取的文本内容，如果无法提取则返回None
        """
        ext = os.path.splitext(file_path)[1].lower()
        
        try:
            if ext in ['.txt', '.md', '.json']:
                return file_content.decode('utf-8')
            elif ext in ['.docx']:
                # 使用python-docx库提取文本（如果已安装）
                try:
                    import docx
                    doc = docx.Document(file_path)
                    return '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                except ImportError:
                    self._logger.warning("未安装python-docx库，无法提取Word文档文本")
                    return None
            elif ext in ['.pdf']:
                # 使用PyPDF2库提取文本（如果已安装）
                try:
                    import PyPDF2
                    with open(file_path, 'rb') as f:
                        reader = PyPDF2.PdfReader(f)
                        text = ''
                        for page in reader.pages:
                            text += page.extract_text() + '\n'
                        return text
                except ImportError:
                    self._logger.warning("未安装PyPDF2库，无法提取PDF文档文本")
                    return None
            else:
                self._logger.warning(f"不支持的文件格式: {ext}")
                return None
        except Exception as e:
            self._logger.error(f"提取文本失败: {e}")
            return None
    
    # ========== 查询操作 ==========
    
    def get_document(self, document_id: str) -> Optional[SolutionDocument]:
        """
        根据ID获取文档
        
        :param document_id: 文档ID
        :return: 文档对象，如果不存在返回None
        """
        document = self._documents.get(document_id)
        if document:
            self._logger.info(f"查询文档: {document_id}")
        else:
            self._logger.warning(f"文档不存在: {document_id}")
        return document
    
    def list_documents(self, filter_condition: Optional[DocumentQueryFilter] = None) -> List[SolutionDocument]:
        """
        查询文档列表
        
        :param filter_condition: 查询过滤条件
        :return: 符合条件的文档列表
        """
        documents = list(self._documents.values())
        
        if filter_condition:
            if filter_condition.document_type:
                documents = [d for d in documents if d.document_type == filter_condition.document_type]
            if filter_condition.file_format:
                documents = [d for d in documents if d.format == filter_condition.file_format]
            if filter_condition.created_by:
                documents = [d for d in documents if d.created_by == filter_condition.created_by]
            if filter_condition.keyword:
                keyword = filter_condition.keyword.lower()
                documents = [
                    d for d in documents
                    if (d.file_name and keyword in d.file_name.lower()) or
                       (d.text_content and keyword in d.text_content.lower()) or
                       (d.description and keyword in d.description.lower())
                ]
        
        self._logger.info(f"查询到 {len(documents)} 个文档")
        return documents
    
    def get_all_documents(self) -> List[SolutionDocument]:
        """
        获取所有文档
        
        :return: 所有文档列表
        """
        return list(self._documents.values())
    
    def document_exists(self, document_id: str) -> bool:
        """
        检查文档是否存在
        
        :param document_id: 文档ID
        :return: 如果存在返回True
        """
        return document_id in self._documents
    
    # ========== 更新操作 ==========
    
    def update_document(self, document_id: str, **kwargs) -> SolutionDocument:
        """
        更新文档属性
        
        :param document_id: 文档ID
        :param kwargs: 要更新的属性
        :return: 更新后的文档对象
        :raises ValueError: 如果文档不存在
        """
        document = self.get_document(document_id)
        if not document:
            raise ValueError(f"文档不存在: {document_id}")
        
        # 使用 Pydantic 的 model_copy 更新属性
        update_data = document.model_dump()
        update_data.update(kwargs)
        update_data['updated_at'] = datetime.now()
        
        updated_document = SolutionDocument(**update_data)
        self._documents[document_id] = updated_document
        
        self._logger.info(f"更新文档: {document_id}")
        return updated_document
    
    def update_document_content(self, document_id: str, text_content: str) -> SolutionDocument:
        """
        更新文档内容
        
        :param document_id: 文档ID
        :param text_content: 新的文本内容
        :return: 更新后的文档对象
        """
        document = self.get_document(document_id)
        if not document:
            raise ValueError(f"文档不存在: {document_id}")
        
        document.text_content = text_content
        document.updated_at = datetime.now()
        
        self._logger.info(f"更新文档内容: {document_id}")
        return document
    
    def update_document_version(self, document_id: str, new_version: str) -> SolutionDocument:
        """
        更新文档版本
        
        :param document_id: 文档ID
        :param new_version: 新版本号
        :return: 更新后的文档对象
        """
        return self.update_document(document_id, version=new_version)
    
    # ========== 删除操作 ==========
    
    def delete_document(self, document_id: str) -> bool:
        """
        删除文档
        
        :param document_id: 文档ID
        :return: 如果删除成功返回True
        :raises ValueError: 如果文档不存在
        """
        if document_id not in self._documents:
            raise ValueError(f"文档不存在: {document_id}")
        
        del self._documents[document_id]
        self._logger.info(f"删除文档: {document_id}")
        return True
    
    # ========== 存储操作 ==========
    
    def save_document(self, document_id: str) -> str:
        """
        保存文档到存储介质
        
        :param document_id: 文档ID
        :return: 保存的文件路径
        """
        document = self.get_document(document_id)
        if not document:
            raise ValueError(f"文档不存在: {document_id}")
        
        file_path = os.path.join(self._storage_path, f"{document_id}.json")
        
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(document.to_json())
        
        self._logger.info(f"保存文档: {document_id} -> {file_path}")
        return file_path
    
    def save_all_documents(self) -> List[str]:
        """
        保存所有文档
        
        :return: 保存的文件路径列表
        """
        saved_paths = []
        for document_id in self._documents:
            path = self.save_document(document_id)
            saved_paths.append(path)
        
        self._logger.info(f"保存了 {len(saved_paths)} 个文档")
        return saved_paths
    
    def load_document(self, file_path: str) -> SolutionDocument:
        """
        从文件加载文档
        
        :param file_path: 文件路径
        :return: 加载的文档对象
        """
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        document = SolutionDocument(**data)
        self._documents[document.document_id] = document
        
        self._logger.info(f"加载文档: {document.document_id} <- {file_path}")
        return document
    
    def load_all_documents(self) -> int:
        """
        从存储目录加载所有文档
        
        :return: 加载的文档数量
        """
        count = 0
        for file_name in os.listdir(self._storage_path):
            if file_name.endswith('.json'):
                file_path = os.path.join(self._storage_path, file_name)
                try:
                    self.load_document(file_path)
                    count += 1
                except Exception as e:
                    self._logger.error(f"加载文档失败: {file_path}, 错误: {e}")
        
        self._logger.info(f"从存储加载了 {count} 个文档")
        return count
    
    # ========== 文本转换操作 ==========
    
    def to_plain_text(self, document_id: str) -> Optional[str]:
        """
        将文档转换为纯文本格式
        
        :param document_id: 文档ID
        :return: 纯文本内容，如果无法转换返回None
        """
        document = self.get_document(document_id)
        if not document:
            raise ValueError(f"文档不存在: {document_id}")
        
        # 如果已有文本内容，直接返回
        if document.text_content:
            return document.text_content
        
        # 如果有文件内容，尝试提取文本
        if document.file_content:
            # 这里简化处理，实际应根据文件格式提取
            self._logger.warning(f"文档 {document_id} 没有预提取的文本内容")
            return None
        
        self._logger.warning(f"文档 {document_id} 没有可转换的内容")
        return None
    
    def to_plain_text_with_metadata(self, document_id: str) -> Optional[str]:
        """
        将文档转换为带元数据的纯文本格式
        
        :param document_id: 文档ID
        :return: 包含元数据的纯文本内容
        """
        document = self.get_document(document_id)
        if not document:
            raise ValueError(f"文档不存在: {document_id}")
        
        text_parts = []
        
        # 添加元数据头
        text_parts.append(f"文档ID: {document.document_id}")
        text_parts.append(f"文件名: {document.file_name}")
        text_parts.append(f"版本: {document.version}")
        text_parts.append(f"类型: {document.document_type}")
        if document.description:
            text_parts.append(f"描述: {document.description}")
        text_parts.append("=" * 50)
        
        # 添加文本内容
        if document.text_content:
            text_parts.append(document.text_content)
        else:
            text_parts.append("[无文本内容]")
        
        return '\n'.join(text_parts)
    
    # ========== 统计操作 ==========
    
    def get_document_count(self) -> int:
        """
        获取文档总数
        
        :return: 文档数量
        """
        return len(self._documents)
    
    def get_statistics(self) -> Dict[str, Any]:
        """
        获取文档统计信息
        
        :return: 统计信息字典
        """
        documents = self.get_all_documents()
        
        type_counts = {}
        format_counts = {}
        total_size = 0
        
        for doc in documents:
            type_counts[doc.document_type.value] = type_counts.get(doc.document_type.value, 0) + 1
            if doc.format:
                format_counts[doc.format] = format_counts.get(doc.format, 0) + 1
            if doc.size:
                total_size += doc.size
        
        return {
            "total_count": len(documents),
            "type_distribution": type_counts,
            "format_distribution": format_counts,
            "total_size_bytes": total_size,
            "total_size_mb": round(total_size / (1024 * 1024), 2) if total_size > 0 else 0
        }